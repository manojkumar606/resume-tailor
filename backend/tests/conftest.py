"""Test fixtures.

Tests run against an in-memory SQLite database so the suite needs no network
and no running Postgres. That is a deliberate trade-off: it keeps CI fast, but
it will not catch Postgres-specific schema problems. Migrations are still
verified against real Postgres via `alembic upgrade head`.
"""

import os
import tempfile

# Must be set before app.core.config is imported anywhere — Settings is cached
# with lru_cache and reads the environment once, at first import.
os.environ.setdefault("DATABASE_URL", "sqlite+pysqlite:///:memory:")
os.environ.setdefault("SECRET_KEY", "test-secret-key")
# Keep uploads out of the real backend/uploads directory.
os.environ.setdefault(
    "LOCAL_STORAGE_DIR", tempfile.mkdtemp(prefix="resume-tailor-tests-")
)

import pytest  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from sqlalchemy import create_engine, select  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402
from sqlalchemy.pool import StaticPool  # noqa: E402

import re  # noqa: E402

from app.core.db import get_db  # noqa: E402
from app.main import app  # noqa: E402
from app.models import Base  # noqa: E402
from app.api.deps import get_mailer  # noqa: E402
from app.services.llm import get_llm_provider  # noqa: E402


class FakeEmailProvider:
    """Captures messages instead of sending them.

    Tests read the verification link straight out of the captured body, so the
    real signup → email → confirm flow is exercised end to end without any
    shortcut that writes is_verified directly.
    """

    def __init__(self):
        self.sent: list[dict] = []

    def send(self, *, to: str, subject: str, html: str, text: str) -> None:
        self.sent.append({"to": to, "subject": subject, "html": html, "text": text})

    def last_code_for(self, email: str) -> str:
        for message in reversed(self.sent):
            if message["to"] == email.strip().lower():
                match = re.search(r"\b(\d{6})\b", message["text"])
                if match:
                    return match.group(1)
        raise AssertionError(f"no code email captured for {email}")


@pytest.fixture
def db_session():
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,  # one shared connection, so :memory: persists
    )
    Base.metadata.create_all(engine)
    session = sessionmaker(bind=engine, autoflush=False, autocommit=False)()
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


@pytest.fixture
def mailbox():
    return FakeEmailProvider()


@pytest.fixture
def client(db_session, mailbox):
    app.dependency_overrides[get_db] = lambda: db_session
    app.dependency_overrides[get_mailer] = lambda: mailbox
    with TestClient(app) as c:
        yield c
    app.dependency_overrides.clear()


@pytest.fixture
def user_payload():
    return {
        "email": "Alice@Example.com",
        "password": "correct-horse-battery",
        "full_name": "Alice Example",
    }


@pytest.fixture
def make_user(client, mailbox):
    """Create a user and return their Authorization header.

    Goes through the real two-step flow — sign up, read the emailed code, submit
    it — rather than writing is_verified directly, so the fixture cannot pass
    while the actual code path is broken.

    Two distinct users in one test is the setup that catches cross-tenant leaks,
    so this is a factory rather than a single fixture.
    """

    def _make(email: str = "user@example.com", password: str = "a-good-password"):
        r = client.post(
            "/api/v1/auth/signup", json={"email": email, "password": password}
        )
        assert r.status_code == 202, r.text

        r = client.post(
            "/api/v1/auth/verify-code",
            json={"email": email, "code": mailbox.last_code_for(email)},
        )
        assert r.status_code == 200, r.text
        assert r.json()["user"]["is_verified"] is True
        return {"Authorization": f"Bearer {r.json()['access_token']}"}

    return _make


@pytest.fixture
def sign_in(client, mailbox):
    """Log an existing user in through both steps, returning fresh headers."""

    def _sign_in(email: str, password: str = "a-good-password"):
        r = client.post(
            "/api/v1/auth/login", json={"email": email, "password": password}
        )
        assert r.status_code == 202, r.text

        r = client.post(
            "/api/v1/auth/verify-code",
            json={"email": email, "code": mailbox.last_code_for(email)},
        )
        assert r.status_code == 200, r.text
        return {"Authorization": f"Bearer {r.json()['access_token']}"}

    return _sign_in


@pytest.fixture
def unverify(db_session):
    """Force an account back to unverified while keeping its token valid.

    A token can now only be obtained by consuming a code, so the verification
    gate is unreachable through the API. It is kept as defence in depth, which
    means testing it requires building the state directly.
    """

    def _unverify(email: str):
        from app.models.user import User

        user = db_session.scalar(
            select(User).where(User.email == email.strip().lower())
        )
        assert user is not None, f"no such user: {email}"
        user.is_verified = False
        db_session.commit()

    return _unverify


@pytest.fixture
def auth_headers(make_user):
    return make_user()


@pytest.fixture
def docx_bytes():
    """A minimal but realistic .docx, including a table (skills are usually
    laid out in one, and table text is easy to drop during extraction)."""
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("Manoj Kumar")
    doc.add_paragraph("Senior Data Analyst")
    doc.add_paragraph("Built reporting pipelines in SQL and Python.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "SQL, Python, dbt"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


DEFAULT_LLM_PAYLOAD = {
    "tailored_resume": (
        "Manoj Kumar\n"
        "EXPERIENCE\n"
        "- Built reporting pipelines in SQL and Python for analytics teams.\n"
        "SKILLS\n"
        "- SQL, Python, dbt, Snowflake\n"
    ),
    "match_score": 78,
    "missing_keywords": ["Snowflake administration"],
    "changes": ["Surfaced dbt experience in the skills section"],
}


DEFAULT_IMPORT_PAYLOAD = {
    "title": "Backend Engineer",
    "company": "Northwind Labs",
    "location": "Remote, India",
    "apply_by": "2026-09-30",
    "description": (
        "We are hiring a Backend Engineer to build and scale our Python API "
        "platform. You will design REST APIs with FastAPI and model data in "
        "PostgreSQL."
    ),
    "confidence": "high",
}

# Content is irrelevant — the fake provider never decodes it — but a real PNG
# signature keeps the fixture honest about what is being uploaded.
PNG_BYTES = b"\x89PNG\r\n\x1a\n" + b"\x00" * 64


class FakeLLM:
    """Deterministic stand-in for a real provider.

    Records its calls so tests can assert on what was actually sent to the
    model, and never touches the network.
    """

    model_name = "fake-model-1"

    def __init__(self, payload=None, error=None):
        self.payload = payload if payload is not None else dict(DEFAULT_LLM_PAYLOAD)
        self.image_payload = dict(DEFAULT_IMPORT_PAYLOAD)
        self.error = error
        self.calls: list[dict] = []

    def generate_json(self, *, system: str, prompt: str) -> dict:
        self.calls.append({"system": system, "prompt": prompt})
        if self.error is not None:
            raise self.error
        return self.payload

    def generate_json_from_images(
        self, *, system: str, prompt: str, images: list[tuple[bytes, str]]
    ) -> dict:
        self.calls.append({"system": system, "prompt": prompt, "images": images})
        if self.error is not None:
            raise self.error
        return self.image_payload


@pytest.fixture
def fake_llm(client):
    fake = FakeLLM()
    app.dependency_overrides[get_llm_provider] = lambda: fake
    yield fake
    app.dependency_overrides.pop(get_llm_provider, None)


CRON_SECRET = "test-cron-secret"


@pytest.fixture
def cron(monkeypatch):
    """Enable the reminder endpoint and return the header that authenticates it."""
    from app.core import config

    monkeypatch.setattr(config.settings, "CRON_SECRET", CRON_SECRET)
    return {"X-Cron-Secret": CRON_SECRET}


@pytest.fixture
def job_payload():
    return {
        "title": "Senior Data Analyst",
        "company": "Acme Corp",
        "location": "Remote",
        "description": (
            "We are looking for a Senior Data Analyst with 5+ years of "
            "experience in SQL, Python and data visualisation. Experience "
            "with dbt and Snowflake strongly preferred."
        ),
    }
