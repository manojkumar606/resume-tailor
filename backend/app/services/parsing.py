"""Extract plain text from uploaded resume files."""

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}
MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5 MB — a resume is never legitimately larger


class UnsupportedFileType(Exception):
    pass


class ParseError(Exception):
    pass


def _from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))

    parts = [p.text for p in doc.paragraphs]
    # Skills and contact details are very often laid out in tables, which
    # doc.paragraphs does not reach. Missing them would silently degrade
    # every tailoring result.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(p for p in parts if p.strip())


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(p.strip() for p in pages if p.strip())


def extract_text(filename: str, data: bytes) -> str:
    """Return the plain text of a resume upload.

    Raises UnsupportedFileType for unknown extensions and ParseError when the
    file is the right type but unreadable (corrupt, or a scanned PDF with no
    text layer).
    """
    if len(data) > MAX_UPLOAD_BYTES:
        raise ParseError(
            f"File is larger than {MAX_UPLOAD_BYTES // (1024 * 1024)} MB"
        )

    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type {suffix!r}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        if suffix == ".docx":
            text = _from_docx(data)
        elif suffix == ".pdf":
            text = _from_pdf(data)
        else:
            text = data.decode("utf-8", errors="replace")
    except (UnsupportedFileType, ParseError):
        raise
    except Exception as exc:
        raise ParseError(f"Could not read the file: {exc}") from exc

    text = text.strip()
    if not text:
        # Most common cause: a PDF that is a scan with no embedded text layer.
        raise ParseError(
            "No text could be extracted. If this is a scanned PDF, upload a "
            ".docx or a text-based PDF instead."
        )
    return text
