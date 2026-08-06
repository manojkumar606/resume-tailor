"""Render tailored resume text into a .docx.

The original prototype copied formatting by matching the Nth tailored line to
the Nth source paragraph. That silently mangles the document whenever the
rewrite changes the line count, which it almost always does. This builds a
clean document from the text instead, inferring headings from shape.
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# A heading is a short line that is not a sentence and not a bullet: either
# ALL CAPS, or Title Case with no trailing period.
_BULLET = re.compile(r"^\s*[-•*•]\s+")


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 60 or _BULLET.match(stripped):
        return False
    if stripped.endswith((".", ",", ";", ":")) and not stripped.isupper():
        return False
    letters = [c for c in stripped if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def build_resume_docx(tailored_text: str, candidate_name: str | None = None) -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    lines = [ln.rstrip() for ln in tailored_text.splitlines()]
    # The model is asked not to emit these, but strip them defensively so a
    # stray separator never shows up in the user's download.
    lines = [ln for ln in lines if ln.strip() != "---SECTION---"]

    first_content_written = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Treat the very first line as the candidate's name.
        if not first_content_written:
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            first_content_written = True
            continue

        if _is_heading(stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11.5)
        elif _BULLET.match(stripped):
            doc.add_paragraph(_BULLET.sub("", stripped), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
