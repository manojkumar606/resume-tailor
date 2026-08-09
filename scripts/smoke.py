#!/usr/bin/env python3
"""End-to-end smoke test: drives the real user journey against a live API.

    # local
    backend/.venv/bin/python scripts/smoke.py

    # deployed
    backend/.venv/bin/python scripts/smoke.py https://resume-tailor-api.onrender.com

Every login now requires a code emailed to the user, and no script can read an
inbox, so the run has two parts:

  * Always — health, and the auth surface: signup issues no token, a wrong
    password is refused, bad codes are rejected, resend does not leak whether an
    address exists.

  * The full journey — upload, storage, tailoring, download — only when
    SMOKE_TOKEN holds a valid access token. Sign in through the browser, then
    copy it out of devtools:

        localStorage.getItem('resume-tailor.token')

        SMOKE_TOKEN=eyJhbGci... backend/.venv/bin/python scripts/smoke.py <url>

Signup leaves a throwaway account behind; harmless, but worth knowing before
running it against something you care about.
"""

from __future__ import annotations

import io
import os
import sys
import time
import uuid

import httpx
from docx import Document

DEFAULT_BASE = "http://127.0.0.1:8000"

# Render free instances sleep after inactivity and can take ~60s to wake, so
# the first request needs a far longer timeout than a warm one would.
TIMEOUT = httpx.Timeout(180.0, connect=90.0)

JOB = {
    "title": "Backend Engineer (Python)",
    "company": "Northwind Labs",
    "location": "Remote, India",
    "description": (
        "We are hiring a Backend Engineer to build and scale our Python API "
        "platform. You will design REST APIs with FastAPI, model data in "
        "PostgreSQL, and own services end to end. Requirements: 2+ years of "
        "Python, strong SQL, experience with REST API design, familiarity "
        "with Docker and CI/CD pipelines. Nice to have: Kubernetes, AWS, and "
        "async task queues such as Celery."
    ),
}


def sample_resume() -> bytes:
    """A small but realistic .docx, including a table — skills are usually laid
    out in one, and table text is easy to lose during extraction."""
    doc = Document()
    for line in (
        "Sample Candidate",
        "sample@example.com | Hyderabad, India",
        "SUMMARY",
        "Software engineer with 3 years building backend services.",
        "EXPERIENCE",
        "Software Engineer, Example Corp - 2022 to present",
        "- Built Python services that process assessment data.",
        "- Automated reporting, cutting manual effort by 6 hours a week.",
        "EDUCATION",
        "B.Tech, Computer Science - 2022",
    ):
        doc.add_paragraph(line)

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL, PostgreSQL, FastAPI, Git"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def step(text: str) -> None:
    print(f"-> {text}", flush=True)


def check_auth_surface(client: httpx.Client, failures: list[str]) -> None:
    step("health")
    started = time.time()
    r = client.get("/health/ready")
    r.raise_for_status()
    print(f"   {r.json()}  ({time.time() - started:.1f}s — slow means a cold start)")

    step("signup issues a code, not a token")
    throwaway = f"smoke-{uuid.uuid4().hex[:8]}@example.com"
    r = client.post("/auth/signup", json={"email": throwaway, "password": "smoke-pw-123"})
    if r.status_code != 202:
        failures.append(f"signup returned {r.status_code}, expected 202")
        print(f"   FAILED {r.status_code}: {r.text[:300]}")
        return
    body = r.json()
    # The whole point of the change: a password alone never yields access.
    if "access_token" in body:
        failures.append("signup handed out a token before the code was confirmed")
    print(f"   {throwaway} -> {body.get('status')}, expires in {body.get('expires_in_minutes')}min")

    step("login with a wrong password is refused")
    r = client.post(
        "/auth/login", json={"email": throwaway, "password": "not-the-password"}
    )
    if r.status_code != 401:
        failures.append(f"wrong password gave {r.status_code}, expected 401")

    step("login for an unknown address is indistinguishable")
    r = client.post(
        "/auth/login", json={"email": "nobody@example.com", "password": "whatever-123"}
    )
    if r.status_code != 401:
        failures.append(f"unknown email gave {r.status_code}, expected 401")
    elif r.json().get("detail") != "Incorrect email or password":
        failures.append("unknown email and wrong password give different messages")

    step("a bogus code is rejected")
    r = client.post("/auth/verify-code", json={"email": throwaway, "code": "000000"})
    if r.status_code != 400:
        failures.append(f"bogus code gave {r.status_code}, expected 400")

    step("resend does not reveal whether an address exists")
    known = client.post("/auth/resend-code", json={"email": throwaway})
    unknown = client.post("/auth/resend-code", json={"email": "nobody@example.com"})
    # A 429 on the known address is fine — the signup email started a cooldown.
    if unknown.status_code != 202:
        failures.append(f"resend for an unknown address gave {unknown.status_code}")
    if known.status_code == 202 and known.json() != unknown.json():
        failures.append("resend responses differ between known and unknown addresses")

    step("protected routes need a token")
    if client.get("/resumes").status_code != 401:
        failures.append("/resumes was reachable without a token")


def run_journey(client: httpx.Client, token: str, failures: list[str]) -> None:
    headers = {"Authorization": f"Bearer {token}"}

    step("token identifies a user")
    r = client.get("/auth/me", headers=headers)
    if r.status_code != 200:
        print(f"   FAILED {r.status_code}: {r.text[:200]}")
        failures.append("SMOKE_TOKEN is not valid — sign in again and copy a fresh one")
        return
    print(f"   {r.json()['email']}, verified={r.json()['is_verified']}")

    step("upload resume")
    r = client.post(
        "/resumes",
        headers=headers,
        files={"file": ("resume.docx", sample_resume(), "application/octet-stream")},
    )
    r.raise_for_status()
    resume = r.json()
    print(f"   parsed {len(resume['parsed_text'])} chars")
    if "PostgreSQL" not in resume["parsed_text"]:
        failures.append("table text was not extracted from the .docx")

    step("download the original back (proves storage round-trips)")
    r = client.get(f"/resumes/{resume['id']}/download", headers=headers)
    r.raise_for_status()
    print(f"   {len(r.content)} bytes")
    if r.content[:2] != b"PK":
        failures.append("stored resume did not come back as a valid .docx")

    step("create job")
    r = client.post("/jobs", headers=headers, json=JOB)
    r.raise_for_status()
    job = r.json()

    step("tailor (live LLM call)")
    started = time.time()
    r = client.post("/tailorings", headers=headers, json={"job_id": job["id"]})
    if r.status_code != 201:
        print(f"   FAILED {r.status_code}: {r.text[:400]}")
        failures.append(f"tailoring returned {r.status_code}")
    else:
        t = r.json()
        print(f"   {t['status']} in {time.time() - started:.1f}s")
        print(f"   match_score      : {t['match_score']}")
        print(f"   missing_keywords : {t['missing_keywords']}")

        step("download tailored .docx")
        r = client.get(f"/tailorings/{t['id']}/download", headers=headers)
        r.raise_for_status()
        print(f"   {len(r.content)} bytes")
        if r.content[:2] != b"PK":
            failures.append("tailored file was not a valid .docx")

    step("cleanup")
    client.delete(f"/jobs/{job['id']}", headers=headers)
    client.delete(f"/resumes/{resume['id']}", headers=headers)
    print("   removed the job and resume created by this run")


def main(argv: list[str]) -> int:
    origin = (argv[1] if len(argv) > 1 else DEFAULT_BASE).rstrip("/")
    base = f"{origin}/api/v1"
    print(f"Target: {base}\n")

    client = httpx.Client(base_url=base, timeout=TIMEOUT, follow_redirects=True)
    failures: list[str] = []

    check_auth_surface(client, failures)

    token = os.environ.get("SMOKE_TOKEN")
    if not token:
        print(
            "\nSkipping the full journey: no SMOKE_TOKEN. Sign in through the\n"
            "browser and copy localStorage.getItem('resume-tailor.token')."
        )
    else:
        print()
        run_journey(client, token, failures)

    print()
    if failures:
        print("SMOKE TEST FAILED")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("SMOKE TEST PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
